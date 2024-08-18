import os
import asyncio
import requests
import logging
from telegram import Bot
from telegram.constants import ParseMode
from telegram.error import TelegramError
from pymongo import MongoClient
import random
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import subprocess
import tempfile
import math
from deep_translator import GoogleTranslator

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize the translator
translator = GoogleTranslator()

# Read environment variables
mongo_uri = os.getenv('MONGO_CONNECTION_STRING')
BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
DEFAULT_CHANNEL = os.getenv('TELEGRAM_CHANNEL_USERNAME')
TEMPLATE_URL = os.getenv('TEMPLATE_URL')

# Debug logging for environment variables
logger.info(f"Bot Token: {BOT_TOKEN}")
logger.info(f"Telegram Channel: {DEFAULT_CHANNEL}")
logger.info(f"Mongo URI: {mongo_uri}")

# Initialize MongoDB client and Telegram bot
client = MongoClient(mongo_uri)
bot = Bot(token=BOT_TOKEN)

# MongoDB collection for tracking selected collections
SELECTED_COLLECTIONS_DB = 'QuizTracker'
SELECTED_COLLECTIONS_COLLECTION = 'SelectedCollections'

def fetch_collections(database_name):
    db = client[database_name]
    return db.list_collection_names()

def fetch_questions_from_collection(database_name, collection_name, num_questions):
    db = client[database_name]
    collection = db[collection_name]
    questions = collection.aggregate([{ '$sample': { 'size': num_questions } }])
    return list(questions)

def get_correct_option_index(answer_key):
    option_mapping = {'a': 0, 'b': 1, 'c': 2, 'd': 3}
    return option_mapping.get(answer_key.lower(), None)

def get_quiz_day():
    db = client['QuizDays']
    collection = db['Days']
    today = datetime.now().date()
    today_datetime = datetime.combine(today, datetime.min.time())

    day_record = collection.find_one({'date': today_datetime})
    
    if day_record:
        return day_record['day']
    else:
        last_day_record = collection.find_one(sort=[('date', -1)])
        new_day = 1 if not last_day_record else last_day_record['day'] + 1
        
        collection.insert_one({'date': today_datetime, 'day': new_day})
        return new_day

def get_quiz_number(collection_name):
    db = client['QuizCounters']
    collection = db['Counters']
    counter_record = collection.find_one({'collection_name': collection_name})
    
    if counter_record:
        new_count = counter_record['count'] + 1
        collection.update_one({'collection_name': collection_name}, {'$set': {'count': new_count}})
        return new_count
    else:
        collection.insert_one({'collection_name': collection_name, 'count': 1})
        return 1

def get_overall_quiz_number():
    db = client['QuizCounters']
    collection = db['OverallCounter']
    counter_record = collection.find_one({'counter_name': 'overall_quiz'})
    
    if counter_record:
        new_count = counter_record['count'] + 1
        collection.update_one({'counter_name': 'overall_quiz'}, {'$set': {'count': new_count}})
        return new_count
    else:
        collection.insert_one({'counter_name': 'overall_quiz', 'count': 1})
        return 1

async def send_intro_message(collection_name, num_questions, quiz_number, overall_quiz_number):
    day = get_quiz_day()
    intro_message = (
        f"🎯 *આજની કવિઝ - Day {day} - {collection_name} Quiz {quiz_number}* 🎯\n\n"
        f"📚 વિષય: *{collection_name}*\n"
        f"🔢 પ્રશ્નોની સંખ્યા: *{num_questions}*\n"
        f"🔢 કવિઝ નંબર: *{overall_quiz_number}*\n\n"
        f"🕐 અમારા ટેલીગ્રામ ચેનલમાં દરરોજ બપોરે *1 વાગ્યે* અને રાત્રે *9 વાગ્યે* "
        f"*{num_questions}* પ્રશ્નોની કવિઝ મુકવામાં આવે છે.\n\n"
        f"🔗 *Join* : @CurrentAdda\n\n"
        f"🏆 તૈયાર રહો! કવિઝ શરૂ થવાની તૈયારીમાં છે... 🚀"
    )
    
    try:
        await bot.send_message(
            chat_id=DEFAULT_CHANNEL,
            text=intro_message,
            parse_mode=ParseMode.MARKDOWN
        )
        logger.info("Intro message sent successfully")
    except TelegramError as e:
        logger.error(f"Error sending intro message: {e}")

async def send_quiz_to_channel(question, options, correct_option_index, explanation):
    question_text = f"{question}\n[@CurrentAdda]"
    
    if explanation is None or (isinstance(explanation, float) and math.isnan(explanation)):
        explanation = "@CurrentAdda"
    
    try:
        await bot.send_poll(
            chat_id=DEFAULT_CHANNEL,
            question=question_text,
            options=options,
            type='quiz',
            correct_option_id=correct_option_index,
            explanation=explanation,
            is_anonymous=True,
            allows_multiple_answers=False,
        )
        logger.info(f"Quiz sent successfully: {question}")
    except TelegramError as e:
        logger.error(f"Error sending quiz: {e}")

def download_template(url):
    download_url = url.replace('/edit?usp=sharing', '/export?format=docx')
    try:
        response = requests.get(download_url)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except requests.exceptions.RequestException as e:
        logger.error(f"Error downloading template: {e}")
        raise

def update_document_with_content(doc_io, intro_message, questions, collection_name, quiz_number):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx_file:
        temp_docx_file.write(doc_io.read())
        temp_docx_path = temp_docx_file.name
    
    doc = Document(temp_docx_path)
    
    content_start = None
    content_end = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        if '<<START_CONTENT>>' in paragraph.text:
            content_start = i
        elif '<<END_CONTENT>>' in paragraph.text:
            content_end = i
            break
    
    if content_start is not None and content_end is not None:
        # Clear existing paragraphs between placeholders
        for i in range(content_end - 1, content_start, -1):
            doc._element.body.remove(doc.paragraphs[i]._element)
        
        # Insert intro message
        intro_para = doc.paragraphs[content_start]
        intro_para.text = intro_message
        intro_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        intro_para.style.font.size = Pt(14)
        intro_para.style.font.bold = True
        intro_para.style.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue color
        
        # Insert questions
        for q in questions:
            question_text = q.get('Question', 'No question text')
            question_paragraph = doc.add_paragraph(f"Q: {question_text}", style='Normal')
            question_paragraph.insert_paragraph_before()
            question_paragraph.style.font.size = Pt(12)
            
            options = [
                f"A) {q.get('Option A', 'No option')}",
                f"B) {q.get('Option B', 'No option')}",
                f"C) {q.get('Option C', 'No option')}",
                f"D) {q.get('Option D', 'No option')}"
            ]
            for option in options:
                option_paragraph = doc.add_paragraph(option, style='Normal')
                option_paragraph.style.font.size = Pt(10)
            
            answer = f"Answer: {q.get('Answer', 'Not provided')}"
            answer_paragraph = doc.add_paragraph(answer, style='Normal')
            answer_paragraph.style.font.size = Pt(10)
            answer_paragraph.style.font.bold = True
            
            doc.add_paragraph()  # Add a blank line between questions
    else:
        logger.warning("Could not find both <<START_CONTENT>> and <<END_CONTENT>> placeholders in the document.")

    updated_doc_path = os.path.join(tempfile.gettempdir(), f'{collection_name} Quiz {quiz_number}.docx')
    doc.save(updated_doc_path)
    
    return updated_doc_path

def convert_docx_to_pdf(docx_file, pdf_path):
    try:
        output_dir = os.path.dirname(pdf_path)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_file],
            check=True, capture_output=True, text=True
        )
        logger.info(f"LibreOffice conversion output: {result.stdout}")
        logger.error(f"LibreOffice conversion error output: {result.stderr}")
        
        pdf_temp_path = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_file))[0] + '.pdf')
        if os.path.exists(pdf_temp_path):
            os.rename(pdf_temp_path, pdf_path)
            logger.info(f"Successfully converted DOCX to PDF: {pdf_path}")
        else:
            logger.error(f"Conversion failed, PDF not found at {pdf_temp_path}")
    except subprocess.CalledProcessError as e:
        logger.error(f"Error converting DOCX to PDF: {e}")

async def send_pdf_to_channel(pdf_path, message_text):
    try:
        with open(pdf_path, 'rb') as pdf_file:
            await bot.send_document(
                chat_id=DEFAULT_CHANNEL,
                document=pdf_file,
                caption=message_text,
                parse_mode=ParseMode.MARKDOWN
            )
            logger.info(f"PDF sent successfully to channel.")
    except TelegramError as e:
        logger.error(f"Error sending PDF: {e}")

async def main():
    database_name = 'YourDatabaseName'  # Replace with your database name
    collection_name = 'YourCollectionName'  # Replace with your collection name
    num_questions = 5  # Adjust as needed
    quiz_number = get_quiz_number(collection_name)
    overall_quiz_number = get_overall_quiz_number()

    # Fetch questions
    questions = fetch_questions_from_collection(database_name, collection_name, num_questions)

    # Create the intro message
    intro_message = (
        f"🎯 *Today's Quiz - Day {get_quiz_day()} - {collection_name} Quiz {quiz_number}* 🎯\n\n"
        f"📚 Topic: *{collection_name}*\n"
        f"🔢 Number of Questions: *{num_questions}*\n"
        f"🔢 Quiz Number: *{overall_quiz_number}*\n\n"
        f"🕐 Daily quizzes are posted in our Telegram channel at *1 PM* and *9 PM* "
        f"with *{num_questions}* questions.\n\n"
        f"🔗 *Join* : @CurrentAdda\n\n"
        f"🏆 Get ready! The quiz is about to start... 🚀"
    )

    # Send the intro message
    await send_intro_message(collection_name, num_questions, quiz_number, overall_quiz_number)

    # Prepare document and convert to PDF
    doc_io = download_template(TEMPLATE_URL)
    doc_path = update_document_with_content(doc_io, intro_message, questions, collection_name, quiz_number)
    pdf_path = os.path.join(tempfile.gettempdir(), f'{collection_name} Quiz {quiz_number}.pdf')
    convert_docx_to_pdf(doc_path, pdf_path)

    # Send the PDF to the channel
    await send_pdf_to_channel(pdf_path, intro_message)

if __name__ == "__main__":
    asyncio.run(main())
